#!/usr/bin/env python3
"""
TAL DDG (Documentation Generation) Tool
Walks through TAL knowledge graph and generates comprehensive architecture documentation
using LLM analysis.
"""

import json
import networkx as nx
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, asdict
from collections import defaultdict
import logging

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


@dataclass
class ArchitectureComponent:
    """Represents a component in the architecture"""
    name: str
    type: str  # procedure, module, service, etc.
    description: str
    dependencies: List[str]
    call_count: int
    complexity: str  # low, medium, high
    business_capability: str


@dataclass
class ProcessFlow:
    """Represents a process flow through the system"""
    name: str
    entry_point: str
    steps: List[Dict[str, Any]]
    data_flow: List[str]
    decision_points: List[str]


class TALDocumentationGenerator:
    """Main class for generating architecture documentation from TAL knowledge graph"""
    
    def __init__(self, graph_path: Optional[str] = None):
        """
        Initialize the documentation generator
        
        Args:
            graph_path: Path to the knowledge graph file (GraphML or JSON)
        """
        self.graph = None
        self.components = []
        self.process_flows = []
        self.statistics = {}
        
        if graph_path:
            self.load_graph(graph_path)
    
    def load_graph(self, graph_path: str):
        """Load the TAL knowledge graph"""
        logger.info(f"Loading knowledge graph from {graph_path}")
        
        path = Path(graph_path)
        if path.suffix == '.graphml':
            self.graph = nx.read_graphml(graph_path)
        elif path.suffix == '.json':
            with open(graph_path, 'r') as f:
                data = json.load(f)
                
                # NetworkX node_link_graph expects 'links' key for edges
                # Handle common variations: 'edges', 'links', 'relationships'
                edge_key = None
                for key in ['links', 'edges', 'relationships']:
                    if key in data:
                        edge_key = key
                        break
                
                if edge_key is None:
                    raise ValueError("JSON must contain 'links', 'edges', or 'relationships' key for graph edges")
                
                # Ensure 'nodes' key exists
                if 'nodes' not in data:
                    raise ValueError("JSON must contain 'nodes' key")
                
                # Create a clean data structure for NetworkX
                # This avoids issues with modifying the original dict
                clean_data = {
                    'directed': data.get('directed', True),
                    'multigraph': data.get('multigraph', False),
                    'graph': data.get('graph', {}),
                    'nodes': data['nodes'],
                    'links': data[edge_key]  # Always use 'links' for NetworkX
                }
                
                # Handle null directed value
                if clean_data['directed'] is None:
                    clean_data['directed'] = True
                    logger.warning("'directed' key was null, assuming directed graph")
                
                self.graph = nx.node_link_graph(clean_data)
        else:
            raise ValueError(f"Unsupported graph format: {path.suffix}")
        
        logger.info(f"Loaded graph with {self.graph.number_of_nodes()} nodes and {self.graph.number_of_edges()} edges")
    
    def analyze_graph(self):
        """Analyze the knowledge graph and extract architectural information"""
        logger.info("Analyzing knowledge graph...")
        
        # Extract components
        self.components = self._extract_components()
        
        # Identify process flows
        self.process_flows = self._identify_process_flows()
        
        # Calculate statistics
        self.statistics = self._calculate_statistics()
        
        logger.info(f"Found {len(self.components)} components and {len(self.process_flows)} process flows")
    
    def _extract_components(self) -> List[ArchitectureComponent]:
        """Extract architectural components from the graph"""
        components = []
        
        for node, data in self.graph.nodes(data=True):
            # Get node attributes
            node_type = data.get('type', 'unknown')
            
            # Skip system intrinsics or utility nodes
            if node_type in ['intrinsic', 'system', 'utility']:
                continue
            
            # Get dependencies (outgoing edges)
            dependencies = list(self.graph.successors(node))
            
            # Calculate in-degree (how many times this is called)
            call_count = self.graph.in_degree(node)
            
            # Determine complexity based on out-degree and call patterns
            out_degree = self.graph.out_degree(node)
            if out_degree > 10 or call_count > 20:
                complexity = "high"
            elif out_degree > 5 or call_count > 10:
                complexity = "medium"
            else:
                complexity = "low"
            
            component = ArchitectureComponent(
                name=node,
                type=node_type,
                description=data.get('description', ''),
                dependencies=dependencies[:10],  # Limit for readability
                call_count=call_count,
                complexity=complexity,
                business_capability=data.get('business_capability', 'unknown')
            )
            components.append(component)
        
        # Sort by call count (most called first)
        components.sort(key=lambda x: x.call_count, reverse=True)
        return components
    
    def _identify_process_flows(self) -> List[ProcessFlow]:
        """Identify key process flows through the system"""
        flows = []
        
        # Find entry points (nodes with high out-degree and low in-degree)
        entry_points = [
            node for node, in_deg in self.graph.in_degree()
            if in_deg <= 2 and self.graph.out_degree(node) >= 3
        ]
        
        for entry_point in entry_points[:10]:  # Limit to top 10 entry points
            # Perform BFS to trace the flow
            steps = []
            visited = set()
            queue = [(entry_point, 0)]
            
            while queue and len(steps) < 20:  # Limit flow depth
                current, depth = queue.pop(0)
                if current in visited:
                    continue
                
                visited.add(current)
                node_data = self.graph.nodes[current]
                
                steps.append({
                    'node': current,
                    'type': node_data.get('type', 'unknown'),
                    'depth': depth,
                    'description': node_data.get('description', '')
                })
                
                # Add successors to queue
                for successor in self.graph.successors(current):
                    if successor not in visited:
                        queue.append((successor, depth + 1))
            
            if len(steps) >= 3:  # Only include meaningful flows
                flow = ProcessFlow(
                    name=f"Flow from {entry_point}",
                    entry_point=entry_point,
                    steps=steps,
                    data_flow=[s['node'] for s in steps],
                    decision_points=[s['node'] for s in steps if self.graph.out_degree(s['node']) > 2]
                )
                flows.append(flow)
        
        return flows
    
    def _calculate_statistics(self) -> Dict[str, Any]:
        """Calculate various statistics about the codebase"""
        stats = {
            'total_procedures': self.graph.number_of_nodes(),
            'total_calls': self.graph.number_of_edges(),
            'avg_dependencies': sum(d for n, d in self.graph.out_degree()) / max(self.graph.number_of_nodes(), 1),
            'max_dependencies': max(d for n, d in self.graph.out_degree()) if self.graph.number_of_nodes() > 0 else 0,
            'most_called': max(self.graph.in_degree(), key=lambda x: x[1])[0] if self.graph.number_of_nodes() > 0 else None,
            'complexity_distribution': defaultdict(int)
        }
        
        # Calculate complexity distribution
        for comp in self.components:
            stats['complexity_distribution'][comp.complexity] += 1
        
        # Identify strongly connected components (circular dependencies)
        strongly_connected = list(nx.strongly_connected_components(self.graph))
        stats['circular_dependency_groups'] = len([c for c in strongly_connected if len(c) > 1])
        
        return stats
    
    def call_llm(self, prompt: str, system_prompt: str = "", temperature: float = 0.7, max_tokens: int = 2000) -> str:
        """
        Call LLM using OpenAI-compatible API
        
        Args:
            prompt: The user prompt
            system_prompt: System prompt for context
            temperature: Sampling temperature
            max_tokens: Maximum tokens in response
        
        Returns:
            LLM response text (always returns a string, never None)
        """
        # This will be replaced with actual OpenAI API call at work
        # For now, return a placeholder
        
        try:
            # PLACEHOLDER - Replace with actual API call
            # Example implementation for OpenAI:
            """
            import openai
            
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            content = response.choices[0].message.content
            # Safety check: ensure we always return a string
            if content is None:
                return "Error: LLM returned None"
            return content
            """
            
            # Example implementation for Anthropic Claude:
            """
            import anthropic
            
            client = anthropic.Anthropic(api_key="your-api-key")
            
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=max_tokens,
                temperature=temperature,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Extract text from response
            content = message.content[0].text if message.content else None
            # Safety check: ensure we always return a string
            if content is None:
                return "Error: LLM returned None"
            return content
            """
            
            logger.warning("LLM call placeholder - implement actual API call")
            return f"[LLM Response Placeholder]\nPrompt: {prompt[:100]}..."
            
        except Exception as e:
            logger.error(f"Error calling LLM: {e}")
            return f"Error: {str(e)}"
    
    def generate_diagram_graphviz(self, diagram_type: str, context: Dict[str, Any]) -> str:
        """
        Generate Graphviz DOT diagram syntax using LLM
        
        Args:
            diagram_type: Type of diagram (architecture, component, process_flow, microservices)
            context: Context data for diagram generation
        
        Returns:
            Graphviz DOT syntax
        """
        
        if diagram_type == "architecture":
            prompt = f"""
Generate a Graphviz DOT diagram showing the high-level architecture of this TAL payment system.

COMPONENT TYPES:
{json.dumps(context.get('component_types', {}), indent=2)}

TOP COMPONENTS BY BUSINESS CAPABILITY:
{json.dumps(context.get('capabilities', {}), indent=2)}

Create a Graphviz digraph showing:
1. Main business capability layers (use subgraphs/clusters)
2. Key components in each layer
3. Major data flows between layers (directed edges)
4. Use different shapes: box for services, cylinder for databases, oval for processes
5. Color code by layer: blue tones for payment, yellow for compliance, green for ledger

IMPORTANT: Return ONLY valid Graphviz DOT code wrapped in ```dot and ```.
Use this structure:
```dot
digraph Architecture {{
    rankdir=TB;
    node [shape=box, style=filled];
    
    subgraph cluster_payment {{
        label="Payment Layer";
        color=lightblue;
        // components here
    }}
    
    // edges showing data flow
}}
```
"""
        
        elif diagram_type == "process_flow":
            prompt = f"""
Generate a Graphviz DOT flowchart showing this process flow:

FLOW: {context.get('flow_name')}
STEPS: {json.dumps(context.get('steps', []), indent=2)}
DECISION POINTS: {json.dumps(context.get('decision_points', []), indent=2)}

Create a Graphviz digraph showing:
1. Process start (oval/ellipse shape)
2. Each step (box shape)
3. Decision points (diamond shape)
4. Alternative paths with labeled edges
5. End points (oval/ellipse shape)
6. Color code: green for start, red for errors, blue for normal flow

IMPORTANT: Return ONLY valid Graphviz DOT code wrapped in ```dot and ```.
Use this structure:
```dot
digraph ProcessFlow {{
    rankdir=TB;
    node [shape=box, style=filled, fillcolor=lightblue];
    
    start [label="Start", shape=ellipse, fillcolor=lightgreen];
    step1 [label="Step 1"];
    decision1 [label="Decision?", shape=diamond, fillcolor=lightyellow];
    
    start -> step1;
    step1 -> decision1;
    decision1 -> step2 [label="Yes"];
    decision1 -> error [label="No"];
}}
```
"""
        
        elif diagram_type == "component":
            prompt = f"""
Generate a Graphviz DOT component diagram for this component:

COMPONENT: {context.get('name')}
TYPE: {context.get('type')}
DEPENDENCIES: {json.dumps(context.get('dependencies', []), indent=2)}
CALLED BY: {json.dumps(context.get('callers', []), indent=2)}

Create a Graphviz digraph showing:
1. The main component (highlighted/different color)
2. Components that call it (callers) - connect with arrows TO the main component
3. Its dependencies (components it calls) - connect with arrows FROM the main component
4. Use different colors: orange for main component, light blue for callers, light green for dependencies

IMPORTANT: Return ONLY valid Graphviz DOT code wrapped in ```dot and ```.
Use this structure:
```dot
digraph ComponentDependencies {{
    rankdir=LR;
    node [shape=box, style=filled];
    
    // Main component
    "{context.get('name')}" [fillcolor=orange];
    
    // Callers
    // caller1 [fillcolor=lightblue];
    
    // Dependencies  
    // dep1 [fillcolor=lightgreen];
    
    // Edges
    // caller1 -> "{context.get('name')}";
    // "{context.get('name')}" -> dep1;
}}
```
"""
        
        elif diagram_type == "microservices":
            prompt = f"""
Generate a Graphviz DOT diagram showing the proposed microservices architecture:

IDENTIFIED SERVICES:
{json.dumps(context.get('services', []), indent=2)}

Create a Graphviz digraph showing:
1. Each proposed microservice as a cluster/subgraph
2. Key components within each service
3. REST API calls between services (solid arrows)
4. Event-driven communication (dashed arrows)
5. Databases (cylinder shape)
6. Color code each service cluster differently

IMPORTANT: Return ONLY valid Graphviz DOT code wrapped in ```dot and ```.
Use this structure:
```dot
digraph MicroservicesArchitecture {{
    rankdir=TB;
    node [shape=box, style=filled];
    
    subgraph cluster_service1 {{
        label="Payment Service";
        color=lightblue;
        style=filled;
        fillcolor=aliceblue;
        
        payment_api [label="Payment API"];
        // other components
    }}
    
    subgraph cluster_service2 {{
        label="Compliance Service";
        color=lightyellow;
        style=filled;
        fillcolor=lightyellow;
        
        compliance_api [label="Compliance API"];
    }}
    
    // API calls
    payment_api -> compliance_api [label="REST"];
    
    // Events
    payment_api -> event_bus [style=dashed, label="Event"];
}}
```
"""
        
        else:
            return f'digraph {{ "Diagram type not supported: {diagram_type}" }}'
        
        system_prompt = """You are a technical architect generating Graphviz DOT diagrams. 
Return ONLY valid Graphviz DOT syntax wrapped in ```dot code blocks. 
Use clear, concise labels. Follow proper DOT syntax with semicolons and proper clustering.
Use appropriate shapes: box, ellipse, diamond, cylinder, component.
Use colors to distinguish different types of nodes."""
        
        response = self.call_llm(prompt, system_prompt, temperature=0.3, max_tokens=1500)
        
        # Safety check: ensure response is a string
        if response is None:
            logger.warning("LLM returned None, returning empty diagram")
            return f'digraph {{ "No diagram generated - LLM returned None" }}'
        
        if not isinstance(response, str):
            logger.warning(f"LLM returned non-string type: {type(response)}")
            response = str(response)
        
        # Extract DOT code from response
        if "```dot" in response:
            start = response.find("```dot") + len("```dot")
            end = response.find("```", start)
            if end > start:
                return response[start:end].strip()
        elif "```" in response:
            # Try to find any code block
            start = response.find("```") + 3
            end = response.find("```", start)
            if end > start:
                code = response[start:end].strip()
                # Check if it starts with digraph
                if code.strip().startswith('digraph'):
                    return code
        
        # Return as-is if no code blocks found
        return response.strip()
    
    def generate_architecture_overview(self) -> Dict[str, Any]:
        """Generate high-level architecture overview using LLM with diagrams"""
        logger.info("Generating architecture overview...")
        
        # Prepare context for LLM
        context = {
            'statistics': self.statistics,
            'top_components': [asdict(c) for c in self.components[:20]],
            'component_types': defaultdict(int)
        }
        
        for comp in self.components:
            context['component_types'][comp.type] += 1
        
        # Group by business capability
        capabilities = defaultdict(list)
        for comp in self.components[:30]:
            capabilities[comp.business_capability].append(comp.name)
        
        prompt = f"""
Based on the following TAL codebase analysis, provide a comprehensive architecture overview:

STATISTICS:
- Total Procedures: {self.statistics['total_procedures']}
- Total Call Relationships: {self.statistics['total_calls']}
- Average Dependencies per Procedure: {self.statistics['avg_dependencies']:.2f}
- Maximum Dependencies: {self.statistics['max_dependencies']}
- Circular Dependency Groups: {self.statistics.get('circular_dependency_groups', 0)}

COMPONENT TYPES:
{json.dumps(dict(context['component_types']), indent=2)}

TOP 20 MOST CRITICAL COMPONENTS:
{json.dumps(context['top_components'][:10], indent=2)}

Please provide:
1. Overall architecture description (2-3 paragraphs)
2. Key architectural patterns observed
3. Main subsystems and their purposes
4. Integration points and dependencies
5. Architectural strengths and weaknesses
"""
        
        system_prompt = """You are a senior software architect analyzing a legacy TAL (Transaction Application Language) 
payment processing system. Provide clear, technical, and actionable insights."""
        
        overview_text = self.call_llm(prompt, system_prompt, temperature=0.3, max_tokens=2000)
        
        # Generate architecture diagram
        logger.info("Generating architecture diagram...")
        diagram_context = {
            'component_types': dict(context['component_types']),
            'capabilities': {k: v[:5] for k, v in capabilities.items()}  # Top 5 per capability
        }
        architecture_diagram = self.generate_diagram_graphviz("architecture", diagram_context)
        
        return {
            'overview': overview_text,
            'diagram': architecture_diagram
        }
    
    def generate_component_documentation(self, component: ArchitectureComponent) -> Dict[str, Any]:
        """Generate detailed documentation for a component with diagram"""
        
        prompt = f"""
Analyze this TAL component and provide detailed documentation:

COMPONENT: {component.name}
TYPE: {component.type}
COMPLEXITY: {component.complexity}
TIMES CALLED: {component.call_count}
DEPENDENCIES: {len(component.dependencies)}
BUSINESS CAPABILITY: {component.business_capability}

KEY DEPENDENCIES: {', '.join(component.dependencies[:5])}

Provide:
1. Purpose and functionality (2-3 sentences)
2. Role in the overall architecture
3. Key dependencies and why they're needed
4. Potential modernization recommendations
"""
        
        system_prompt = "You are documenting legacy TAL payment processing code. Be specific and technical."
        
        documentation = self.call_llm(prompt, system_prompt, temperature=0.5, max_tokens=800)
        
        # Generate component diagram if it has dependencies
        diagram = None
        if len(component.dependencies) > 0:
            # Get callers (reverse lookup)
            callers = [n for n in self.graph.nodes() if component.name in list(self.graph.successors(n))]
            
            diagram_context = {
                'name': component.name,
                'type': component.type,
                'dependencies': component.dependencies[:8],
                'callers': callers[:5]
            }
            diagram = self.generate_diagram_graphviz("component", diagram_context)
        
        return {
            'documentation': documentation,
            'diagram': diagram
        }
    
    def generate_process_flow_documentation(self, flow: ProcessFlow) -> Dict[str, Any]:
        """Generate documentation for a process flow with diagram"""
        
        flow_steps = "\n".join([f"{i+1}. {step['node']} ({step['type']})" 
                                for i, step in enumerate(flow.steps[:10])])
        
        prompt = f"""
Document this process flow from the TAL payment system:

FLOW NAME: {flow.name}
ENTRY POINT: {flow.entry_point}
NUMBER OF STEPS: {len(flow.steps)}
DECISION POINTS: {len(flow.decision_points)}

FLOW SEQUENCE:
{flow_steps}

DECISION POINTS: {', '.join(flow.decision_points[:5])}

Provide:
1. Business process description (what does this flow accomplish?)
2. Step-by-step explanation of the main path
3. Key decision points and their impact
4. Data transformations in this flow
5. Error handling considerations
"""
        
        system_prompt = "You are a business analyst documenting payment processing workflows."
        
        documentation = self.call_llm(prompt, system_prompt, temperature=0.5, max_tokens=1000)
        
        # Generate process flow diagram
        diagram_context = {
            'flow_name': flow.name,
            'steps': [{'name': s['node'], 'type': s['type']} for s in flow.steps[:12]],
            'decision_points': flow.decision_points[:5]
        }
        diagram = self.generate_diagram_graphviz("process_flow", diagram_context)
        
        return {
            'documentation': documentation,
            'diagram': diagram
        }
    
    def identify_microservices_candidates(self) -> Dict[str, Any]:
        """Use LLM to identify potential microservice boundaries with diagram"""
        logger.info("Identifying microservice candidates...")
        
        # Group components by business capability
        capabilities = defaultdict(list)
        for comp in self.components:
            capabilities[comp.business_capability].append(comp)
        
        candidates = []
        service_summary = []
        
        for capability, comps in capabilities.items():
            if len(comps) < 3:  # Skip small groups
                continue
            
            context = {
                'capability': capability,
                'component_count': len(comps),
                'total_calls': sum(c.call_count for c in comps),
                'components': [c.name for c in comps[:10]]
            }
            
            prompt = f"""
Evaluate this group of TAL components as a potential microservice:

BUSINESS CAPABILITY: {capability}
NUMBER OF COMPONENTS: {len(comps)}
TOTAL INCOMING CALLS: {context['total_calls']}

KEY COMPONENTS:
{json.dumps(context['components'], indent=2)}

Assess:
1. Is this a good microservice candidate? (Yes/No)
2. Cohesion level (High/Medium/Low)
3. Suggested service name
4. API boundaries needed
5. Data ownership scope
"""
            
            response = self.call_llm(prompt, "You are a microservices architect.", temperature=0.3)
            
            candidates.append({
                'capability': capability,
                'components': [asdict(c) for c in comps[:10]],  # Convert to dicts for JSON serialization
                'analysis': response
            })
            
            # Collect for diagram
            service_summary.append({
                'name': capability.replace('_', ' ').title(),
                'component_count': len(comps),
                'key_components': [c.name for c in comps[:3]]
            })
        
        # Generate microservices architecture diagram
        logger.info("Generating microservices architecture diagram...")
        diagram_context = {
            'services': service_summary[:8]  # Limit to top 8 for readability
        }
        diagram = self.generate_diagram_graphviz("microservices", diagram_context)
        
        return {
            'candidates': candidates,
            'diagram': diagram
        }
    
    def generate_data_flow_analysis(self) -> str:
        """Analyze data flows through the system"""
        logger.info("Analyzing data flows...")
        
        # Find longest paths (representing major data flows)
        try:
            # For directed graphs, we look at paths from entry points
            entry_points = [n for n in self.graph.nodes() if self.graph.in_degree(n) <= 1]
            
            sample_paths = []
            for entry in entry_points[:5]:
                # Get paths from this entry point
                paths = nx.single_source_shortest_path(self.graph, entry, cutoff=8)
                if paths:
                    longest = max(paths.values(), key=len)
                    if len(longest) > 3:
                        sample_paths.append(longest)
        except:
            sample_paths = []
        
        prompt = f"""
Analyze data flows in this TAL payment processing system:

STATISTICS:
- Total Procedures: {self.statistics['total_procedures']}
- Total Data Flow Paths: {self.statistics['total_calls']}

SAMPLE DATA FLOW PATHS:
{json.dumps([' -> '.join(path) for path in sample_paths[:5]], indent=2)}

TOP DATA TRANSFORMATION POINTS:
{json.dumps([c.name for c in self.components[:10]], indent=2)}

Provide:
1. Main data flow patterns
2. Data transformation stages
3. Critical data processing nodes
4. Data consistency mechanisms needed
5. Recommended data architecture for modernization
"""
        
        return self.call_llm(prompt, "You are a data architect analyzing payment systems.", temperature=0.4)
    
    def generate_modernization_roadmap(self) -> str:
        """Generate modernization recommendations"""
        logger.info("Generating modernization roadmap...")
        
        prompt = f"""
Create a modernization roadmap for this TAL legacy system:

CURRENT STATE:
- Total Procedures: {self.statistics['total_procedures']}
- Average Dependencies: {self.statistics['avg_dependencies']:.2f}
- Circular Dependencies: {self.statistics.get('circular_dependency_groups', 0)}
- High Complexity Components: {self.statistics['complexity_distribution'].get('high', 0)}
- Medium Complexity: {self.statistics['complexity_distribution'].get('medium', 0)}
- Low Complexity: {self.statistics['complexity_distribution'].get('low', 0)}

ARCHITECTURE PATTERNS:
- Component Types Identified: {len(set(c.type for c in self.components))}
- Process Flows Mapped: {len(self.process_flows)}

Provide a phased modernization roadmap:
1. **Phase 1** (Quick Wins - 0-6 months)
2. **Phase 2** (Core Modernization - 6-18 months)
3. **Phase 3** (Complete Migration - 18-36 months)

For each phase include:
- Key activities
- Technologies to adopt
- Risks and mitigation
- Success metrics
"""
        
        return self.call_llm(prompt, "You are a technology transformation consultant.", temperature=0.4, max_tokens=2500)
    
    def generate_documentation_report(self, output_file: str = "tal_architecture_report.json"):
        """Generate complete documentation report with diagrams"""
        logger.info("Generating comprehensive documentation report...")
        
        report = {
            'metadata': {
                'generated_at': str(Path.cwd()),
                'graph_stats': self.statistics
            },
            'architecture_overview': self.generate_architecture_overview(),
            'components': [],
            'process_flows': [],
            'data_flow_analysis': self.generate_data_flow_analysis(),
            'microservices_candidates': self.identify_microservices_candidates(),
            'modernization_roadmap': self.generate_modernization_roadmap()
        }
        
        # Document top 10 critical components
        logger.info("Documenting critical components...")
        for comp in self.components[:10]:
            comp_doc = self.generate_component_documentation(comp)
            report['components'].append({
                'component': asdict(comp),
                'documentation': comp_doc['documentation'],
                'diagram': comp_doc['diagram']
            })
        
        # Document top 5 process flows
        logger.info("Documenting process flows...")
        for flow in self.process_flows[:5]:
            flow_doc = self.generate_process_flow_documentation(flow)
            report['process_flows'].append({
                'flow': asdict(flow),
                'documentation': flow_doc['documentation'],
                'diagram': flow_doc['diagram']
            })
        
        # Save report
        with open(output_file, 'w') as f:
            json.dump(report, f, indent=2)
        
        logger.info(f"Documentation report saved to {output_file}")
        return report
    
    def render_graphviz_to_image(self, dot_code: str, output_path: str, format: str = 'png') -> bool:
        """
        Render Graphviz DOT diagram to image
        
        Tries multiple methods:
        1. graphviz command (dot command)
        2. Python graphviz library
        3. Save as .dot file for manual rendering
        
        Args:
            dot_code: Graphviz DOT syntax
            output_path: Output file path
            format: Output format (png, svg, pdf)
        
        Returns:
            True if successful, False otherwise
        """
        try:
            import subprocess
            import tempfile
            
            # Create temporary dot file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.dot', delete=False) as f:
                f.write(dot_code)
                dot_file = f.name
            
            try:
                # Try using graphviz command line (dot)
                result = subprocess.run(
                    ['dot', f'-T{format}', dot_file, '-o', output_path],
                    capture_output=True,
                    timeout=30
                )
                
                if result.returncode == 0 and Path(output_path).exists():
                    logger.info(f"Rendered Graphviz diagram to {output_path}")
                    return True
                    
            except (subprocess.TimeoutExpired, FileNotFoundError):
                logger.warning("Graphviz (dot command) not available, trying Python library...")
                
                # Try using Python graphviz library
                try:
                    import graphviz
                    source = graphviz.Source(dot_code)
                    source.render(output_path.replace(f'.{format}', ''), format=format, cleanup=True)
                    logger.info(f"Rendered Graphviz diagram using Python library to {output_path}")
                    return True
                except ImportError:
                    logger.warning("Python graphviz library not installed")
                except Exception as e:
                    logger.warning(f"Failed to render with Python library: {e}")
            
            # If both methods failed, save the .dot file for reference
            dot_output = output_path.replace(f'.{format}', '.dot')
            with open(dot_output, 'w') as f:
                f.write(dot_code)
            logger.info(f"Saved Graphviz source to {dot_output} (install graphviz to render)")
            
            return False
            
        except Exception as e:
            logger.warning(f"Failed to render Graphviz diagram: {e}")
            return False
    
    def create_word_document(self, report: Dict[str, Any], output_file: str = "tal_architecture_guide.docx"):
        """Create a Word document from the report with embedded diagrams"""
        logger.info("Creating Word document...")
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('TAL Architecture & Modernization Guide', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Table of Contents placeholder
            doc.add_heading('Table of Contents', 1)
            doc.add_paragraph('[Auto-generated in Word]')
            doc.add_page_break()
            
            # 1. Architecture Overview
            doc.add_heading('1. Architecture Overview', 1)
            overview_data = report.get('architecture_overview', {})
            if isinstance(overview_data, dict):
                doc.add_paragraph(overview_data.get('overview', 'N/A'))
                
                # Add architecture diagram
                if overview_data.get('diagram'):
                    doc.add_heading('1.1 Architecture Diagram', 2)
                    diagram_path = 'architecture_diagram.png'
                    if self.render_graphviz_to_image(overview_data['diagram'], diagram_path):
                        doc.add_picture(diagram_path, width=Inches(6))
                    else:
                        # Include DOT code as code block
                        doc.add_paragraph('Graphviz DOT Diagram Code:', style='Intense Quote')
                        p = doc.add_paragraph(overview_data['diagram'])
                        p.style = 'Code'
            else:
                doc.add_paragraph(str(overview_data))
            
            # Statistics
            doc.add_heading('1.2 System Statistics', 2)
            stats = report['metadata']['graph_stats']
            table = doc.add_table(rows=len(stats), cols=2)
            table.style = 'Light Grid Accent 1'
            for i, (key, value) in enumerate(stats.items()):
                if key != 'complexity_distribution':
                    table.rows[i].cells[0].text = str(key).replace('_', ' ').title()
                    table.rows[i].cells[1].text = str(value)
            
            doc.add_page_break()
            
            # 2. Components
            doc.add_heading('2. Key Components', 1)
            for i, comp_data in enumerate(report['components'][:10]):
                comp = comp_data['component']
                doc.add_heading(f"2.{i+1} {comp['name']}", 2)
                doc.add_paragraph(f"Type: {comp['type']} | Complexity: {comp['complexity']} | Called: {comp['call_count']} times")
                doc.add_paragraph(comp_data.get('documentation', 'N/A'))
                
                # Add component diagram
                if comp_data.get('diagram'):
                    diagram_path = f"component_{i}_diagram.png"
                    if self.render_graphviz_to_image(comp_data['diagram'], diagram_path):
                        doc.add_picture(diagram_path, width=Inches(5))
                    else:
                        doc.add_paragraph('Component Diagram (Graphviz):', style='Intense Quote')
                        p = doc.add_paragraph(comp_data['diagram'][:200] + '...')
                        p.style = 'Code'
            
            doc.add_page_break()
            
            # 3. Process Flows
            doc.add_heading('3. Process Flows', 1)
            for i, flow_data in enumerate(report['process_flows'][:5]):
                flow = flow_data['flow']
                doc.add_heading(f"3.{i+1} {flow['name']}", 2)
                doc.add_paragraph(flow_data.get('documentation', 'N/A'))
                
                # Add process flow diagram
                if flow_data.get('diagram'):
                    diagram_path = f"flow_{i}_diagram.png"
                    if self.render_graphviz_to_image(flow_data['diagram'], diagram_path):
                        doc.add_picture(diagram_path, width=Inches(6))
                    else:
                        doc.add_paragraph('Process Flow Diagram (Graphviz):', style='Intense Quote')
                        p = doc.add_paragraph(flow_data['diagram'][:200] + '...')
                        p.style = 'Code'
            
            doc.add_page_break()
            
            # 4. Data Flow Analysis
            doc.add_heading('4. Data Flow Analysis', 1)
            doc.add_paragraph(report.get('data_flow_analysis', 'N/A'))
            
            doc.add_page_break()
            
            # 5. Microservices Candidates
            doc.add_heading('5. Microservices Candidates', 1)
            microservices_data = report.get('microservices_candidates', {})
            
            # Add microservices diagram
            if isinstance(microservices_data, dict) and microservices_data.get('diagram'):
                doc.add_heading('5.1 Proposed Microservices Architecture', 2)
                diagram_path = 'microservices_diagram.png'
                if self.render_graphviz_to_image(microservices_data['diagram'], diagram_path):
                    doc.add_picture(diagram_path, width=Inches(6.5))
                else:
                    doc.add_paragraph('Microservices Architecture (Graphviz):', style='Intense Quote')
                    p = doc.add_paragraph(microservices_data['diagram'][:300] + '...')
                    p.style = 'Code'
                
                doc.add_heading('5.2 Service Details', 2)
                candidates = microservices_data.get('candidates', [])
            else:
                candidates = microservices_data if isinstance(microservices_data, list) else []
            
            for candidate in candidates[:5]:
                doc.add_heading(f"5.2.{candidate.get('capability', 'Unknown')}", 3)
                doc.add_paragraph(f"Components: {len(candidate.get('components', []))}")
                doc.add_paragraph(candidate.get('analysis', 'N/A'))
            
            doc.add_page_break()
            
            # 6. Modernization Roadmap
            doc.add_heading('6. Modernization Roadmap', 1)
            doc.add_paragraph(report.get('modernization_roadmap', 'N/A'))
            
            # Save document
            doc.save(output_file)
            logger.info(f"Word document saved to {output_file}")
            
            # Note about diagram rendering
            doc_note = """
NOTE: Diagrams are included as images if Graphviz is installed.
To render diagrams manually:
1. Install Graphviz: apt-get install graphviz (Linux) or brew install graphviz (Mac)
2. Render .dot files: dot -Tpng diagram.dot -o diagram.png
Or view online: https://dreampuf.github.io/GraphvizOnline/
"""
            logger.info(doc_note)
            
        except ImportError:
            logger.warning("python-docx not installed. Skipping Word document generation.")
            logger.info("Install with: pip install python-docx")


def main():
    """Main execution function"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate TAL architecture documentation from knowledge graph')
    parser.add_argument('graph_file', help='Path to knowledge graph file (GraphML or JSON)')
    parser.add_argument('--output-json', default='tal_architecture_report.json', help='Output JSON report file')
    parser.add_argument('--output-docx', default='tal_architecture_guide.docx', help='Output Word document file')
    parser.add_argument('--skip-docx', action='store_true', help='Skip Word document generation')
    
    args = parser.parse_args()
    
    # Initialize generator
    generator = TALDocumentationGenerator(args.graph_file)
    
    # Analyze graph
    generator.analyze_graph()
    
    # Generate documentation report
    report = generator.generate_documentation_report(args.output_json)
    
    # Create Word document
    if not args.skip_docx:
        generator.create_word_document(report, args.output_docx)
    
    logger.info("Documentation generation complete!")


if __name__ == '__main__':
    main()
